"""Word Twist - 지문·문제 누적 저장소 (Ollama 포함)"""
import streamlit as st
import json
import os
import re
import uuid
import datetime
import concurrent.futures

import google.generativeai as genai
try:
    from openai import OpenAI
except Exception:
    OpenAI = None
try:
    import anthropic
except Exception:
    anthropic = None

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from io import BytesIO
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

# Ollama 지원
try:
    import ollama
    OLLAMA_AVAILABLE = True
except ImportError:
    OLLAMA_AVAILABLE = False

from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload

SAVE_FILE = "saved_questions.json"
PASSAGE_FILE = "passages.json"
SCOPES = ['https://www.googleapis.com/auth/drive.file']

POS_ROTATION = ["verb", "adjective", "adverb", "conjunction"]
POS_KOR = {"verb": "동사", "adjective": "형용사", "adverb": "부사", "conjunction": "접속사"}

GRAMMAR_CATEGORIES = [
    "주어동사_수일치", "시제", "능동수동", "조동사_가정법", "to부정사_동명사",
    "분사_분사구문", "현재분사_과거분사", "형용사_부사", "비교구문",
    "명사_대명사", "재귀대명사", "관계대명사", "관계부사", "what_which_that_whose",
    "접속사_전치사", "병렬구조", "도치", "강조구문", "어순", "5형식_보어",
]

GRAMMAR_HINTS = {
    "주어동사_수일치": "긴 주어구의 핵을 잘못 잡거나 주어-동사 단복수 어긋나게.",
    "시제": "현재완료/과거완료/미래완료/시제일치 위반.",
    "능동수동": "자동사를 수동으로(is happened), 또는 능동/수동 부적절.",
    "조동사_가정법": "가정법 if절-주절 동사 짝, should/would/could 시제 어긋나게.",
    "to부정사_동명사": "동명사만 받는 동사에 to부정사 (enjoy to do), 그 반대.",
    "분사_분사구문": "분사구문 주어와 본문 주어 불일치, 분사구문 형태 오류.",
    "현재분사_과거분사": "능동의 V-ing vs 수동의 V-ed 혼동(boring vs bored).",
    "형용사_부사": "동사·형용사를 수식하는데 형용사 사용(부사가 와야 하는 자리).",
    "비교구문": "the 비교급, the 비교급 / as 원급 as / than 동사 일치 위반.",
    "명사_대명사": "단수/복수, 대명사 선행사 일치, it/them/its 혼동.",
    "재귀대명사": "재귀 vs 인칭 (himself vs him), 강조용법 vs 재귀용법.",
    "관계대명사": "선행사 격 (who/whom/whose/which/that) 잘못된 격.",
    "관계부사": "where/when/why/how — 선행사가 장소/시간/이유/방법인지 어긋나게.",
    "what_which_that_whose": "선행사 유무에 따른 what vs that, whose vs of which.",
    "접속사_전치사": "because vs because of, while vs during, although vs despite.",
    "병렬구조": "and/or/but 양쪽 형태 불일치 (V-ing & to V).",
    "도치": "Only/Never/Hardly 등 부정어 도치 위반, So/Neither 도치.",
    "강조구문": "It is ~ that 강조구문에서 엉뚱한 자리 강조하거나 형태 오류.",
    "어순": "간접의문문 어순 (의문사+S+V vs 의문사+V+S) 위반.",
    "5형식_보어": "지각·사역동사 보어 형태 (see/hear/feel/make/let + 원형/V-ing/p.p.).",
}

GEMINI_MODELS = ["gemini-2.0-flash", "gemini-1.5-flash-latest", "gemini-2.5-flash", "gemini-flash-latest", "gemini-pro-latest"]
OPENAI_MODELS = ["gpt-4o", "gpt-4o-mini", "gpt-4-turbo", "gpt-3.5-turbo"]
CLAUDE_MODELS = ["claude-opus-4-6", "claude-sonnet-4-6", "claude-haiku-4-5-20251001", "claude-3-5-sonnet-20241022", "claude-3-5-haiku-20241022"]
DEEPSEEK_MODELS = ["deepseek-chat", "deepseek-reasoner"]


# ==================== 지문 관리 ====================
def now_iso():
    return datetime.datetime.now().isoformat()


def load_passages():
    if not os.path.exists(PASSAGE_FILE):
        return []
    try:
        with open(PASSAGE_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return []


def save_passages(passages):
    with open(PASSAGE_FILE, "w", encoding="utf-8") as f:
        json.dump(passages, f, ensure_ascii=False, indent=4)


def add_passage(title, text):
    passages = load_passages()
    for p in passages:
        if p["title"] == title:
            p["text"] = text
            p["updated_at"] = now_iso()
            save_passages(passages)
            return p["id"]
    new_id = str(uuid.uuid4())
    passages.append({
        "id": new_id,
        "title": title,
        "text": text,
        "created_at": now_iso(),
        "updated_at": now_iso(),
    })
    save_passages(passages)
    return new_id


def delete_passage(passage_id):
    passages = load_passages()
    passages = [p for p in passages if p["id"] != passage_id]
    save_passages(passages)


def get_passage_by_id(passage_id):
    for p in load_passages():
        if p["id"] == passage_id:
            return p
    return None


# ==================== 문제 저장/로드 ====================
def load_all_questions():
    if not os.path.exists(SAVE_FILE):
        return []
    try:
        with open(SAVE_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return []


def save_question(data):
    if "_id" not in data:
        data["_id"] = str(uuid.uuid4())
    if "_saved_at" not in data:
        data["_saved_at"] = now_iso()
    fd = load_all_questions()
    fd.append(data)
    with open(SAVE_FILE, "w", encoding="utf-8") as f:
        json.dump(fd, f, ensure_ascii=False, indent=4)


def delete_question_by_id(qid):
    fd = load_all_questions()
    fd = [q for q in fd if q.get("_id") != qid]
    with open(SAVE_FILE, "w", encoding="utf-8") as f:
        json.dump(fd, f, ensure_ascii=False, indent=4)


def delete_questions_for_text(text):
    """이 지문의 모든 문제 삭제"""
    t = text.strip()
    fd = load_all_questions()
    fd = [q for q in fd if q.get("original_text", "").strip() != t]
    with open(SAVE_FILE, "w", encoding="utf-8") as f:
        json.dump(fd, f, ensure_ascii=False, indent=4)


def delete_all_questions():
    with open(SAVE_FILE, "w", encoding="utf-8") as f:
        json.dump([], f, ensure_ascii=False, indent=4)


def questions_for_text(text):
    t = text.strip()
    return [i for i in load_all_questions() if i.get("original_text", "").strip() == t]


def previous_targets(text):
    return [i.get("original_word", "") for i in questions_for_text(text) if i.get("original_word")]


def previous_pos(text):
    return [i.get("answer_pos", "") for i in questions_for_text(text) if i.get("answer_pos")]


def previous_choice_words(text):
    words = set()
    for q in questions_for_text(text):
        qt = q.get("question_text", "")
        for m in re.findall(r"<u>(.*?)</u>", qt, flags=re.IGNORECASE | re.DOTALL):
            w = m.strip()
            if w:
                words.add(w)
    return sorted(words)


def pick_focus_pos(text):
    history = previous_pos(text)
    if not history:
        return POS_ROTATION[0]
    counts = {p: history.count(p) for p in POS_ROTATION}
    mn = min(counts.values())
    cands = [p for p in POS_ROTATION if counts[p] == mn]
    for p in POS_ROTATION:
        if p in cands and history[-1] != p:
            return p
    return cands[0]


def previous_grammar_categories(text):
    cats = []
    for q in questions_for_text(text):
        if q.get("_q_type") == "grammar":
            c = q.get("answer_pos", "")
            if c:
                cats.append(c)
    return cats


def pick_grammar_category(text):
    import random
    history = previous_grammar_categories(text)
    if not history:
        return random.choice(GRAMMAR_CATEGORIES[:8])
    counts = {c: history.count(c) for c in GRAMMAR_CATEGORIES}
    mn = min(counts.values())
    cands = [c for c in GRAMMAR_CATEGORIES if counts[c] == mn]
    last = history[-1] if history else None
    cands_diff = [c for c in cands if c != last]
    pool = cands_diff if cands_diff else cands
    return random.choice(pool)


# ==================== LLM 호출 (Ollama 포함) ====================
def call_llm(provider, model, prompt, api_keys):
    if provider == "gemini":
        genai.configure(api_key=api_keys.get("gemini", ""))
        return genai.GenerativeModel(model).generate_content(prompt).text
    if provider == "openai":
        if OpenAI is None:
            raise RuntimeError("pip install openai 필요")
        c = OpenAI(api_key=api_keys.get("openai", ""))
        r = c.chat.completions.create(model=model, messages=[{"role": "user", "content": prompt}], temperature=0.9)
        return r.choices[0].message.content
    if provider == "anthropic":
        if anthropic is None:
            raise RuntimeError("pip install anthropic 필요")
        c = anthropic.Anthropic(api_key=api_keys.get("anthropic", ""))
        r = c.messages.create(model=model, max_tokens=4096, messages=[{"role": "user", "content": prompt}])
        return r.content[0].text
    if provider == "deepseek":
        if OpenAI is None:
            raise RuntimeError("pip install openai 필요")
        c = OpenAI(api_key=api_keys.get("deepseek", ""), base_url="https://api.deepseek.com")
        r = c.chat.completions.create(model=model, messages=[{"role": "user", "content": prompt}], temperature=0.9)
        return r.choices[0].message.content
    if provider == "ollama":
        if not OLLAMA_AVAILABLE:
            raise RuntimeError("Ollama가 설치되지 않았습니다. `pip install ollama` 실행 후 다시 시도하세요.")
        response = ollama.chat(model=model, messages=[{"role": "user", "content": prompt}])
        return response["message"]["content"]
    raise ValueError("unknown provider: " + str(provider))


# ==================== 프롬프트 빌더 ====================
def build_prompt(text, prev_targets, focus_pos=None, prev_choices=None):
    avoid_part = ""
    if prev_targets:
        avoid_part = "\n[중복 절대 금지] 아래 단어들은 이전 정답으로 사용됨. 다시 쓰지 마라: " + ", ".join(prev_targets)
    if prev_choices:
        avoid_part += "\n[다양성 힌트] 이전 보기 단어: " + ", ".join(prev_choices) + " — 가능하면 이번엔 다른 단어를 골라라."

    if focus_pos:
        pos_kor = POS_KOR[focus_pos]
        pos_rule = "5. 이번 정답 타겟 품사는 반드시 " + pos_kor + " (" + focus_pos + ") 다."
        pos_field = focus_pos
    else:
        pos_rule = "5. 정답(변형할 1개)은 동사/형용사/부사/접속사 중 가장 함정성 높은 단어를 자유롭게 골라라."
        pos_field = "verb 또는 adjective 또는 adverb 또는 conjunction 중 실제 변형한 품사"

    template = """너는 대한민국 고등학교 상위권~수능 수준의 영어 어휘 문제 전문가다.
'문맥상 낱말의 쓰임이 적절하지 않은 것은?' 문제를 만든다.

[난이도] 대한민국 고2~고3 상위권 / 수능·모의고사 수준.

[보기 선정]
1. 5개 보기는 형용사/동사/접속사/부사 중에서만 (명사·전치사·관사 금지).
2. 4개 품사가 골고루 섞이게.
3. 5개 모두 문맥 흐름을 결정짓는 핵심 단어.
4. 보기 단어 앞에 (1)~(5) 번호 + HTML <u>단어</u> 태그.

[정답 변형]
__POS_RULE__
6. 원문을 반의어/정반대 의미 단어로 교체 (품사는 동일).
7. 문법은 자연스럽지만 문맥상 명백히 어긋나야 함.
8. 나머지 4개는 원문 그대로.

[전체 흐름]
9. 변형 1개 외 지문 내용·논리·구조는 원문과 100% 동일.
__AVOID__

[출력] 마크다운 코드블록 없이 순수 JSON만:
{
    "question_text": "(<u>(1) word</u> 형태 5개 포함된 영어 지문 전문)",
    "answer": 1,
    "original_word": "변형 전 원래 단어",
    "modified_word": "변형 후 들어간 단어",
    "answer_pos": "__POS_FIELD__",
    "explanation": "왜 부적절한지 한국어 2~3문장"
}

원문:
\"\"\"
__TEXT__
\"\"\"
"""
    return (template
            .replace("__POS_RULE__", pos_rule)
            .replace("__POS_FIELD__", pos_field)
            .replace("__AVOID__", avoid_part)
            .replace("__TEXT__", text))


def build_grammar_prompt(text, prev_targets, prev_choices=None, focus_category=None):
    avoid_part = ""
    if prev_targets:
        avoid_part = "\n[중복 금지] 이전 정답 어법 포인트: " + ", ".join(prev_targets) + " — 다시 쓰지 마라."
    if prev_choices:
        avoid_part += "\n[다양성 힌트] 이전 보기 부분: " + ", ".join(prev_choices) + " — 가능하면 다른 부분을 골라라."

    cat_block = ""
    if focus_category:
        hint = GRAMMAR_HINTS.get(focus_category, "")
        cat_block = (
            "\n[★ 이번 문제 정답 카테고리는 반드시 '" + focus_category + "' 다 ★]\n"
            "정답(틀린 1개)은 반드시 이 카테고리의 어법 포인트여야 한다. "
            "다른 카테고리는 절대 정답으로 만들지 마라.\n"
            "이 카테고리 함정 패턴 가이드: " + hint + "\n"
            "이 함정을 자연스럽게 지문 안에 녹여서 출제하라.\n"
        )

    template = """너는 대한민국 고등학교 상위권~수능 수준의 영어 어법 문제 전문가다.
'다음 글의 밑줄 친 부분 중, 어법상 틀린 것은?' 문제를 만든다.

[난이도] 대한민국 고2~고3 상위권 / 수능·모의고사 어법 수준.

[★★ 보기 5개 그룹 강제 분배 ★★]
5개 보기는 반드시 아래 5개 그룹에서 **각 그룹당 정확히 1개씩** 뽑아라.
같은 그룹에서 2개 이상 뽑으면 절대 안 된다.

  그룹 A (동사 관련): 시제, 능동수동, 분사, 가정법, 5형식 보어, to부정사 vs 동명사
  그룹 B (관계사·접속사·전치사): 관계대명사, 관계부사, what/which/that/whose, 접속사 vs 전치사
  그룹 C (형용사·부사·비교): 형용사 vs 부사, 비교구문, 도치
  그룹 D (대명사·명사·일치): 명사_대명사, 재귀대명사, 주어동사 수일치
  그룹 E (특수구문): 병렬구조, 강조구문, 간접의문문 어순

[★ 절대 하지 말 것 — 나쁜 예시 ★]
보기가 (1) used to seat, (2) who would go, (3) were held, (4) particularly liked, (5) is chatting
→ 5개 모두 그룹 A(동사). 다양성 0. 이런 출제 절대 금지.

[★ 이렇게 출제 — 좋은 예시 ★]
(1) is held [그룹A·능동수동]
(2) which is [그룹B·관계대명사]
(3) particularly [그룹C·부사]
(4) themselves [그룹D·재귀대명사]
(5) to watch [그룹E·병렬구조 / 또는 to부정사]
→ A, B, C, D, E 각 1개씩 골고루.

[보기 단위 규칙]
- 보기는 단어 1개가 아니라 2~5 단어 구 단위 (예: "to study", "having seen", "which is", "the more important", "themselves").
- 5개 모두 실제 어법 포인트가 되는 부분.
- 보기에 (1)~(5) 번호 + HTML <u>해당 부분</u> 태그.
- 관사(a/an/the)는 보기 대상에서 제외.
__CAT_BLOCK__
[정답 변형]
- 5개 중 정확히 1개만 어법상 틀리게 변형 (의미는 그대로, 어법만 틀리게).
- 나머지 4개는 어법상 완벽히 정확한 원문 그대로.

[전체 흐름]
- 변형 1곳 외 지문 전체는 원문과 100% 동일.
__AVOID__

[출력] 마크다운 코드블록 없이 순수 JSON만:
{
    "question_text": "(<u>(1) 부분</u> 형태 5개 포함된 영어 지문 전문)",
    "answer": 1,
    "original_word": "변형 전 원래 어법 형태",
    "modified_word": "변형 후 어법상 틀린 형태",
    "answer_pos": "정답이 속한 카테고리 정확한 이름 (예: 관계대명사, 분사_분사구문, 수일치 등)",
    "explanation": "왜 어법상 틀린지, 어떻게 고쳐야 하는지 한국어 2~3문장"
}

원문:
\"\"\"
__TEXT__
\"\"\"
"""
    return (template
            .replace("__CAT_BLOCK__", cat_block)
            .replace("__AVOID__", avoid_part)
            .replace("__TEXT__", text))


def build_analysis_prompt(text):
    return """너는 한국 고등학생을 가르치는 영어 강사다.
다음 영어 지문을 문장 단위로 정밀 분석해라.

[작업 단위]
각 문장마다 아래 4가지를 출력한다.

1. original: 영어 원문 그대로 (한 문장씩)

2. annotated: 영어 원문에 어구 단위로 인라인 주석 추가
   - 형식: "어구⟦성분·품사⟧ 어구⟦성분·품사⟧ ..."
   - 문장 성분 약어 사용: S(주어), V(동사), O(목적어), C(주격보어), OC(목적격보어), M(수식어/부사구), Conj(접속사)
   - 품사도 같이 표기: 명사구, 동사, 형용사, 부사, 전치사구, 분사구문, to부정사, 관계절, 종속절 등
   - 예시: "Good thinkers⟦S·명사구⟧ rarely⟦M·부사⟧ limit⟦V·타동사⟧ themselves⟦O·재귀대명사⟧ to a single way⟦M·전치사구⟧ of understanding the world⟦M·전치사구·동명사⟧."
   - 절(clause)이 들어있으면 절 단위로 묶어서 표기. 예: "when Galileo finally got around to doing some empirical studies of gravity⟦M·종속절(시간)⟧"

3. literal: 직독직해 — 영어 어구 순서 그대로 한국어로 끊어서 번역
   - 슬래시(/)로 어구 구분
   - 예시: "좋은 사상가들은 / 거의 ~하지 않는다 / 자신을 가두지 / 한 가지 방식에 / 세상을 이해하는"

4. translation: 자연스러운 한국어 해석 (의역, 한 문장)
   - 예시: "훌륭한 사상가들은 세상을 이해하는 한 가지 방식에만 자신을 가두지 않는다."

[출력] 마크다운 코드블록 없이 순수 JSON만:
{
  "sentences": [
    {
      "original": "...",
      "annotated": "...",
      "literal": "...",
      "translation": "..."
    },
    ...
  ]
}

원문:
\"\"\"
__TEXT__
\"\"\"
""".replace("__TEXT__", text)


def analyze_passage(text, provider, model, api_keys):
    prompt = build_analysis_prompt(text)
    raw = call_llm(provider, model, prompt, api_keys)
    cleaned = raw.strip().replace("```json", "").replace("```", "").strip()
    s = cleaned.find("{")
    e = cleaned.rfind("}")
    if s != -1 and e != -1:
        cleaned = cleaned[s:e + 1]
    return json.loads(cleaned)


def generate_one_raw(text, prev_targets, focus_pos, provider, model, api_keys, q_type="vocab", prev_choices=None, focus_category=None):
    if q_type == "grammar":
        prompt = build_grammar_prompt(text, prev_targets, prev_choices, focus_category=focus_category)
    else:
        prompt = build_prompt(text, prev_targets, focus_pos, prev_choices)
    raw = call_llm(provider, model, prompt, api_keys)
    cleaned = raw.strip().replace("```json", "").replace("```", "").strip()
    s = cleaned.find("{")
    e = cleaned.rfind("}")
    if s != -1 and e != -1:
        cleaned = cleaned[s:e + 1]
    return json.loads(cleaned)


# ==================== 구글 드라이브 ====================
def get_drive_service():
    creds = None
    if os.path.exists('token.json'):
        creds = Credentials.from_authorized_user_file('token.json', SCOPES)
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            if os.path.exists('credentials.json'):
                flow = InstalledAppFlow.from_client_secrets_file('credentials.json', SCOPES)
                creds = flow.run_local_server(port=0)
            else:
                st.error("credentials.json 파일이 필요합니다.")
                return None
        with open('token.json', 'w') as token:
            token.write(creds.to_json())
    return build('drive', 'v3', credentials=creds)


def upload_to_drive(path):
    s = get_drive_service()
    if not s:
        return None
    media = MediaFileUpload(path, mimetype='application/json')
    f = s.files().create(body={'name': os.path.basename(path)}, media_body=media, fields='id').execute()
    return f.get('id')


# ==================== 내보내기 (Word) ====================
def export_questions_to_docx(questions, filename="word_twist_questions.docx"):
    if not DOCX_AVAILABLE:
        st.error("python-docx 라이브러리가 필요합니다. `pip install python-docx` 실행 후 다시 시도하세요.")
        return None
    doc = Document()
    doc.add_heading('Word Twist - 생성 문제', 0)
    for i, q in enumerate(questions, 1):
        doc.add_heading(f'문제 {i}', level=1)
        doc.add_paragraph(q.get("question_text", ""), style='Normal')
        doc.add_paragraph(f'정답: {q.get("answer", "?")}번', style='Intense Quote')
        doc.add_paragraph(f'원래 단어: {q.get("original_word", "?")} → 변형: {q.get("modified_word", "?")}')
        doc.add_paragraph(f'카테고리: {q.get("answer_pos", "?")}')
        doc.add_paragraph(f'해설: {q.get("explanation", "")}')
        doc.add_paragraph('---')
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ==================== Streamlit UI ====================
st.set_page_config(page_title="Word Twist", page_icon="🌀", layout="wide")

CSS = """
<style>
.stApp { background: radial-gradient(1200px 600px at 10% -10%, rgba(168,85,247,0.18), transparent 60%), radial-gradient(1000px 500px at 110% 10%, rgba(236,72,153,0.15), transparent 60%), radial-gradient(900px 600px at 50% 120%, rgba(6,182,212,0.12), transparent 60%), linear-gradient(180deg, #0a0e1a 0%, #0b1020 100%); color: #e5e7eb; }
header[data-testid="stHeader"] { background: transparent; }
#MainMenu, footer { visibility: hidden; }
h1, h2, h3, h4 { color: #e5e7eb !important; letter-spacing: -0.01em; }
.hero-title { font-size: 2.6rem; font-weight: 800; line-height: 1.1; background: linear-gradient(90deg, #a855f7, #ec4899, #06b6d4); -webkit-background-clip: text; background-clip: text; color: transparent; margin: 0.2rem 0 0.4rem 0; }
.hero-sub { color: #9ca3af; font-size: 0.95rem; margin-bottom: 1.4rem; }
.chip { display: inline-block; padding: 4px 10px; border-radius: 999px; font-size: 0.75rem; font-weight: 600; margin-right: 6px; margin-bottom: 4px; border: 1px solid rgba(255,255,255,0.08); background: rgba(255,255,255,0.03); color: #9ca3af; }
.chip-accent { background: linear-gradient(90deg, rgba(168,85,247,0.3), rgba(236,72,153,0.3)); color: #fff; border-color: transparent; }
.chip-success { background: linear-gradient(90deg, rgba(34,197,94,0.3), rgba(168,85,247,0.3)); color: #fff; border-color: transparent; }
.q-card { background: linear-gradient(180deg, rgba(255,255,255,0.05), rgba(255,255,255,0.02)); border: 1px solid rgba(255,255,255,0.08); border-radius: 18px; padding: 1.4rem 1.6rem; line-height: 1.85; font-size: 1.02rem; box-shadow: 0 10px 40px rgba(0,0,0,0.3); }
.q-card u { text-decoration: none; border-bottom: 2px solid #a855f7; padding: 0 2px; color: #fff; font-weight: 600; }
.stTextArea textarea { background: rgba(255,255,255,0.03) !important; color: #e5e7eb !important; border: 1px solid rgba(255,255,255,0.08) !important; border-radius: 12px !important; }
.stButton > button { background: linear-gradient(90deg, #a855f7, #ec4899); color: white; border: none; border-radius: 12px; padding: 0.55rem 1.1rem; font-weight: 600; box-shadow: 0 8px 22px rgba(168,85,247,0.25); transition: transform 0.15s ease, box-shadow 0.15s ease; }
.stButton > button:hover { transform: translateY(-1px); box-shadow: 0 12px 28px rgba(236,72,153,0.35); }
section[data-testid="stSidebar"] { background: rgba(10,14,26,0.7); border-right: 1px solid rgba(255,255,255,0.08); }
.metric-row { display: flex; gap: 14px; flex-wrap: wrap; }
.metric { flex: 1; min-width: 140px; background: rgba(255,255,255,0.04); border: 1px solid rgba(255,255,255,0.08); border-radius: 14px; padding: 0.9rem 1rem; }
.metric .label { color: #9ca3af; font-size: 0.75rem; text-transform: uppercase; letter-spacing: 0.06em; }
.metric .value { font-size: 1.5rem; font-weight: 700; margin-top: 0.2rem; }
.divider { height: 1px; background: rgba(255,255,255,0.08); margin: 1.4rem 0; }
.small-dim { color: #9ca3af; font-size: 0.85rem; }
.loading-overlay { position: fixed; top: 0; left: 0; right: 0; bottom: 0; background: radial-gradient(800px 500px at 50% 50%, rgba(168,85,247,0.25), rgba(10,14,26,0.92) 70%); backdrop-filter: blur(10px); -webkit-backdrop-filter: blur(10px); z-index: 99999; display: flex; align-items: center; justify-content: center; animation: fadeIn 0.25s ease; }
@keyframes fadeIn { from {opacity:0} to {opacity:1} }
.loading-box { text-align: center; padding: 38px 56px; border-radius: 22px; background: linear-gradient(180deg, rgba(255,255,255,0.06), rgba(255,255,255,0.02)); border: 1px solid rgba(168,85,247,0.35); box-shadow: 0 30px 80px rgba(0,0,0,0.5), 0 0 60px rgba(168,85,247,0.25); }
.spinner { width: 84px; height: 84px; border-radius: 50%; margin: 0 auto 24px; background: conic-gradient(from 0deg, #a855f7, #ec4899, #06b6d4, #a855f7); -webkit-mask: radial-gradient(circle 32px at center, transparent 98%, #000 100%); mask: radial-gradient(circle 32px at center, transparent 98%, #000 100%); animation: spin 1.2s linear infinite; filter: drop-shadow(0 0 18px rgba(168,85,247,0.5)); }
@keyframes spin { to { transform: rotate(360deg); } }
.loading-title { color: #fff; font-size: 1.35rem; font-weight: 800; letter-spacing: -0.01em; margin-bottom: 6px; background: linear-gradient(90deg, #c084fc, #ec4899); -webkit-background-clip: text; background-clip: text; color: transparent; }
.loading-sub { color: #cbd5e1; font-size: 0.95rem; }
.loading-dots::after { content: ''; display: inline-block; width: 1em; text-align: left; animation: dots 1.4s steps(4, end) infinite; }
@keyframes dots { 0%{content:''} 25%{content:'.'} 50%{content:'..'} 75%{content:'...'} 100%{content:''} }

.sent-block { background: linear-gradient(180deg, rgba(255,255,255,0.04), rgba(255,255,255,0.015)); border: 1px solid rgba(255,255,255,0.08); border-radius: 16px; padding: 1.2rem 1.4rem; margin-bottom: 1rem; }
.sent-num { display: inline-block; padding: 2px 10px; border-radius: 999px; background: linear-gradient(90deg, #a855f7, #ec4899); color:#fff; font-size: 0.75rem; font-weight: 700; margin-bottom: 10px; }
.sent-original { color:#e5e7eb; font-size: 1rem; line-height:1.7; margin-bottom: 12px; font-weight: 500; }
.sent-row { color:#cbd5e1; font-size: 0.93rem; line-height: 1.85; margin-top: 8px; padding-top: 8px; border-top: 1px solid rgba(255,255,255,0.06); }
.sent-row .lbl { color: #c084fc; font-weight: 700; font-size: 0.72rem; letter-spacing: 0.08em; text-transform: uppercase; margin-right: 8px; }
.sent-row.literal { color: #fbcfe8; font-family: "Pretendard", sans-serif; }
.sent-row.translation { color: #f5f3ff; font-size: 0.97rem; }
.sent-annotated { color:#fff; font-size: 0.97rem; line-height: 2.0; word-spacing: 2px; }
.tag-chip { display: inline-block; padding: 1px 6px; border-radius: 6px; font-size: 0.7rem; font-weight: 700; margin-left: 2px; vertical-align: 1px; letter-spacing: 0.02em; }
.tag-S  { background: rgba(59,130,246,0.25); color: #93c5fd; border: 1px solid rgba(59,130,246,0.4); }
.tag-V  { background: rgba(236,72,153,0.25); color: #f9a8d4; border: 1px solid rgba(236,72,153,0.4); }
.tag-O  { background: rgba(34,197,94,0.25);  color: #86efac; border: 1px solid rgba(34,197,94,0.4); }
.tag-C  { background: rgba(234,179,8,0.25);  color: #fde68a; border: 1px solid rgba(234,179,8,0.4); }
.tag-OC { background: rgba(249,115,22,0.25); color: #fdba74; border: 1px solid rgba(249,115,22,0.4); }
.tag-M  { background: rgba(148,163,184,0.18); color: #cbd5e1; border: 1px solid rgba(148,163,184,0.3); }
</style>
"""
st.markdown(CSS, unsafe_allow_html=True)


# ---------- 세션 상태 초기화 ----------
if "current_text" not in st.session_state:
    st.session_state.current_text = ""
if "last_results" not in st.session_state:
    st.session_state.last_results = []
if "selected_passage_id" not in st.session_state:
    st.session_state.selected_passage_id = None
if "passage_select_label" not in st.session_state:
    st.session_state.passage_select_label = "(직접 입력)"


# ==================== 사이드바 ====================
with st.sidebar:
    st.markdown("### 🔌 모델 / API 키")
    provider_options = ["gemini", "openai", "anthropic", "deepseek"]
    if OLLAMA_AVAILABLE:
        provider_options.append("ollama")
    provider = st.radio("Provider", provider_options,
        format_func=lambda x: {"gemini": "Gemini", "openai": "GPT", "anthropic": "Claude", "deepseek": "DeepSeek", "ollama": "Ollama (로컬)"}[x],
        horizontal=True)

    if provider == "gemini":
        model = st.selectbox("모델", GEMINI_MODELS, index=0)
        gemini_key = st.text_input("Gemini API Key", type="password", value=os.environ.get("GEMINI_API_KEY", ""))
        st.session_state["gemini_api_key"] = gemini_key
    elif provider == "openai":
        model = st.selectbox("모델", OPENAI_MODELS, index=0)
        openai_key = st.text_input("OpenAI API Key", type="password", value=os.environ.get("OPENAI_API_KEY", ""))
        st.session_state["openai_api_key"] = openai_key
    elif provider == "deepseek":
        model = st.selectbox("모델", DEEPSEEK_MODELS, index=0)
        deepseek_key = st.text_input("DeepSeek API Key", type="password", value=os.environ.get("DEEPSEEK_API_KEY", ""))
        st.session_state["deepseek_api_key"] = deepseek_key
    elif provider == "anthropic":
        model = st.selectbox("모델", CLAUDE_MODELS, index=0)
        anthropic_key = st.text_input("Anthropic API Key", type="password", value=os.environ.get("ANTHROPIC_API_KEY", ""))
        st.session_state["anthropic_api_key"] = anthropic_key
    elif provider == "ollama":
        # Ollama 설치된 모델 목록 가져오기
        try:
            models = ollama.list()['models']
            model_names = [m['name'] for m in models]
            if not model_names:
                st.warning("설치된 Ollama 모델이 없습니다. 터미널에서 `ollama pull llama3.1:8b` 실행 후 다시 시도하세요.")
                model_names = ["llama3.1:8b"]
            model = st.selectbox("모델", model_names, index=0)
        except Exception:
            st.error("Ollama 서버에 연결할 수 없습니다. `ollama serve` 실행 여부를 확인하세요.")
            model = st.text_input("모델명 (직접 입력)", "llama3.1:8b")
        st.info("로컬 Ollama 사용: API 키 불필요, 무제한 생성 가능")

    st.markdown("---")
    st.markdown("### 🎯 출제 옵션")
    q_type_label = st.radio("문제 유형", ["어휘 (문맥상 부적절한 단어)", "어법 (어법상 틀린 것)"], index=0)
    q_type = "grammar" if q_type_label.startswith("어법") else "vocab"
    pos_mode = st.radio("정답 품사 결정 (어휘만)", ["자유", "로테이션", "고정"], index=0, horizontal=True)
    manual_pos = "verb"
    auto_pos = (pos_mode == "로테이션")
    free_pos = (pos_mode == "자유")
    if pos_mode == "고정":
        manual_pos = st.selectbox("정답 품사 고정", POS_ROTATION, format_func=lambda x: POS_KOR[x])
    batch_n = st.number_input("한 번에 만들 문제 수", 1, 10, 1)
    parallel_mode = st.checkbox("병렬 생성 (배치 시 빠름)", value=True)

    st.markdown("---")
    st.markdown("### 📚 지문 보관함")

    passages = load_passages()
    label_to_id = {"(직접 입력)": None}
    for p in passages:
        label = p["title"] + "  [" + p["updated_at"][:10] + "]"
        label_to_id[label] = p["id"]
    passage_options = list(label_to_id.keys())

    current_label = st.session_state.passage_select_label
    if current_label not in passage_options:
        current_label = "(직접 입력)"
    current_idx = passage_options.index(current_label)

    new_label = st.selectbox(
        "저장된 지문 불러오기",
        passage_options,
        index=current_idx,
        key="passage_selector",
    )

    if new_label != st.session_state.passage_select_label:
        st.session_state.passage_select_label = new_label
        target_id = label_to_id[new_label]
        if target_id is None:
            st.session_state.selected_passage_id = None
            st.session_state.current_text = ""
            st.session_state["text_area"] = ""
        else:
            p = get_passage_by_id(target_id)
            if p:
                st.session_state.selected_passage_id = target_id
                st.session_state.current_text = p["text"]
                st.session_state["text_area"] = p["text"]
        st.session_state.last_results = []
        st.rerun()

    new_title = st.text_input("새 지문 제목", placeholder="예: 2025 수능특강 3강")
    if st.button("💾 현재 지문 저장"):
        if not new_title.strip():
            st.warning("제목을 입력해주세요.")
        elif not st.session_state.current_text.strip():
            st.warning("저장할 지문이 없습니다.")
        else:
            pid = add_passage(new_title.strip(), st.session_state.current_text)
            st.session_state.selected_passage_id = pid
            saved_p = get_passage_by_id(pid)
            if saved_p:
                st.session_state.passage_select_label = saved_p["title"] + "  [" + saved_p["updated_at"][:10] + "]"
            st.success("'" + new_title + "' 지문이 저장되었습니다.")
            st.rerun()

    if st.session_state.selected_passage_id:
        if st.button("🗑️ 선택 지문 삭제"):
            delete_passage(st.session_state.selected_passage_id)
            st.session_state.selected_passage_id = None
            st.session_state.current_text = ""
            st.session_state.last_results = []
            st.session_state.passage_select_label = "(직접 입력)"
            st.session_state["text_area"] = ""
            st.success("지문이 삭제되었습니다.")
            st.rerun()

    st.markdown("---")
    st.markdown("### 🧹 문제 관리")
    if st.button("🗑️ 전체 문제 초기화 (모든 지문)", use_container_width=True):
        delete_all_questions()
        st.success("모든 문제가 삭제되었습니다.")
        st.rerun()


# ==================== 메인 영역 ====================
st.markdown("<div class='hero-title'>Word Twist</div>", unsafe_allow_html=True)
st.markdown("<div class='hero-sub'>한 지문 · 무한히 비틀기 — 어휘 + 어법 통합 출제기 (Ollama 지원)</div>", unsafe_allow_html=True)

if st.session_state.selected_passage_id:
    cur_p = get_passage_by_id(st.session_state.selected_passage_id)
    if cur_p:
        st.markdown(
            "<span class='chip chip-success'>✓ 지문 불러옴: " + cur_p["title"] + "</span>",
            unsafe_allow_html=True,
        )

if "text_area" not in st.session_state:
    st.session_state["text_area"] = st.session_state.current_text

user_input = st.text_area(
    "영어 지문",
    height=220,
    placeholder="여기에 영어 지문을 붙여 넣으세요...",
    key="text_area",
)

if user_input.strip() != st.session_state.current_text.strip():
    st.session_state.current_text = user_input
    st.session_state.last_results = []
    if "analysis_result" in st.session_state:
        del st.session_state["analysis_result"]
    if st.session_state.selected_passage_id:
        st.session_state.selected_passage_id = None
        st.session_state.passage_select_label = "(직접 입력)"

tab_q, tab_a = st.tabs(["🌀 문제 생성", "📖 지문 분석"])


# =============== TAB Q: 문제 생성 ===============
with tab_q:
    if user_input.strip():
        existing = questions_for_text(user_input)
        used_targets = [q.get("original_word", "?") for q in existing]
        if q_type == "grammar":
            pos_label = "어법"
        elif free_pos:
            pos_label = "자유"
        elif auto_pos:
            pos_label = POS_KOR.get(pick_focus_pos(user_input), "?")
        else:
            pos_label = POS_KOR.get(manual_pos, "?")

        metric_html = (
            "<div class='metric-row'>"
            "<div class='metric'><div class='label'>이 지문 누적</div><div class='value'>" + str(len(existing)) + "개</div></div>"
            "<div class='metric'><div class='label'>다음 정답</div><div class='value'>" + pos_label + "</div></div>"
            "<div class='metric'><div class='label'>회피 단어 수</div><div class='value'>" + str(len(used_targets)) + "</div></div>"
            "</div>"
        )
        st.markdown(metric_html, unsafe_allow_html=True)
        if used_targets:
            st.markdown("<div class='small-dim' style='margin-top:0.6rem'>이전 정답:</div>", unsafe_allow_html=True)
            chips = "".join("<span class='chip'>" + w + "</span>" for w in used_targets)
            st.markdown(chips, unsafe_allow_html=True)

    st.markdown("<div class='divider'></div>", unsafe_allow_html=True)

    c1, c2, c3 = st.columns([1, 1, 2])
    with c1:
        btn_one = st.button("🌀 새 문제 만들기")
    with c2:
        btn_batch = st.button("⚡ " + str(int(batch_n)) + "개 한 번에", key="btn_batch_q")


# ---------- 문제 생성 로직 ----------
def make_one(text):
    prev_t = previous_targets(text)
    prev_c = previous_choice_words(text)
    if free_pos:
        focus = None
    elif auto_pos:
        focus = pick_focus_pos(text)
    else:
        focus = manual_pos
    grammar_cat = None
    if q_type == "grammar":
        grammar_cat = pick_grammar_category(text)
    api_keys = {
        "gemini": st.session_state.get("gemini_api_key", ""),
        "openai": st.session_state.get("openai_api_key", ""),
        "anthropic": st.session_state.get("anthropic_api_key", ""),
        "deepseek": st.session_state.get("deepseek_api_key", ""),
    }
    last_err = None
    for attempt in range(4):
        try:
            result = generate_one_raw(text, prev_t, focus, provider, model, api_keys,
                                       q_type=q_type, prev_choices=prev_c,
                                       focus_category=grammar_cat)
        except Exception as e:
            last_err = e
            continue
        ow = (result.get("original_word") or "").strip().lower()
        if not ow:
            last_err = "original_word 누락"
            continue
        if ow in {t.lower() for t in prev_t}:
            last_err = "중복 타겟: " + ow
            continue
        result["original_text"] = text.strip()
        result["_provider"] = provider
        result["_model"] = model
        result["_q_type"] = q_type
        if st.session_state.selected_passage_id:
            result["passage_id"] = st.session_state.selected_passage_id
        save_question(result)
        st.session_state.last_results.append(result)
        return result
    raise RuntimeError("생성 실패: " + str(last_err))


def render_loading(placeholder, current, total, provider_name, mode_label=""):
    sub = provider_name.upper() + " · " + str(current) + " / " + str(total)
    if mode_label:
        sub = provider_name.upper() + " · " + mode_label
    hint = "<div style='margin-top:14px; color:#94a3b8; font-size:0.78rem'>한 문제당 보통 8~15초 걸려요</div>"
    overlay_html = (
        "<div class='loading-overlay'><div class='loading-box'>"
        "<div class='spinner'></div>"
        "<div class='loading-title'>문제를 만들고 있어요</div>"
        "<div class='loading-sub'>" + sub + " <span class='loading-dots'></span></div>"
        + hint + "</div></div>"
    )
    placeholder.markdown(overlay_html, unsafe_allow_html=True)


def _make_one_threadsafe(text, prev_targets_snapshot, focus, provider_name, model_name, api_keys,
                          q_type_snapshot, prev_choices_snapshot, focus_category_snapshot=None):
    return generate_one_raw(text, prev_targets_snapshot, focus, provider_name, model_name, api_keys,
                             q_type=q_type_snapshot, prev_choices=prev_choices_snapshot,
                             focus_category=focus_category_snapshot)


with tab_q:
    if btn_one or btn_batch:
        if not user_input.strip():
            st.warning("먼저 영어 지문을 입력해주세요.")
        else:
            n = int(batch_n) if btn_batch else 1
            loader = st.empty()
            ok = 0
            errors = []

            api_keys_snapshot = {
                "gemini": st.session_state.get("gemini_api_key", ""),
                "openai": st.session_state.get("openai_api_key", ""),
                "anthropic": st.session_state.get("anthropic_api_key", ""),
                "deepseek": st.session_state.get("deepseek_api_key", ""),
            }

            if n > 1 and parallel_mode:
                render_loading(loader, 0, n, provider, mode_label=str(n) + "개 동시 생성 중")
                prev_snapshot = previous_targets(user_input)
                prev_choices_snap = previous_choice_words(user_input)

                tasks_focus = []
                if free_pos or q_type == "grammar":
                    tasks_focus = [None] * n
                elif auto_pos:
                    focus_default = pick_focus_pos(user_input)
                    tasks_focus = [POS_ROTATION[(POS_ROTATION.index(focus_default) + i) % len(POS_ROTATION)] for i in range(n)]
                else:
                    tasks_focus = [manual_pos] * n

                tasks_grammar_cat = []
                if q_type == "grammar":
                    grammar_history = previous_grammar_categories(user_input)
                    counts = {c: grammar_history.count(c) for c in GRAMMAR_CATEGORIES}
                    sorted_cats = sorted(GRAMMAR_CATEGORIES, key=lambda c: counts[c])
                    tasks_grammar_cat = sorted_cats[:n]
                else:
                    tasks_grammar_cat = [None] * n

                results = []
                with concurrent.futures.ThreadPoolExecutor(max_workers=min(n, 8)) as ex:
                    futures = {
                        ex.submit(_make_one_threadsafe, user_input, prev_snapshot, tasks_focus[i],
                                  provider, model, api_keys_snapshot, q_type, prev_choices_snap,
                                  tasks_grammar_cat[i]): i for i in range(n)
                    }
                    for f in concurrent.futures.as_completed(futures):
                        try:
                            results.append(f.result())
                        except Exception as e:
                            errors.append(str(e))

                seen = {t.lower() for t in prev_snapshot}
                for r in results:
                    ow = (r.get("original_word") or "").strip().lower()
                    if not ow:
                        errors.append("original_word 누락")
                        continue
                    if ow in seen:
                        errors.append("중복 정답으로 제외: " + ow)
                        continue
                    seen.add(ow)
                    r["original_text"] = user_input.strip()
                    r["_provider"] = provider
                    r["_model"] = model
                    r["_q_type"] = q_type
                    if st.session_state.selected_passage_id:
                        r["passage_id"] = st.session_state.selected_passage_id
                    save_question(r)
                    st.session_state.last_results.append(r)
                    ok += 1
            else:
                for i in range(n):
                    render_loading(loader, i + 1, n, provider)
                    try:
                        make_one(user_input)
                        ok += 1
                    except Exception as e:
                        errors.append(str(e))

            loader.empty()
            if ok:
                st.success(str(ok) + "개 문제 생성 완료")
            if errors:
                with st.expander("⚠️ 실패 로그"):
                    for e in errors:
                        st.write("- " + e)


def render_question_card(idx, r, show_delete_button=False):
    prov = str(r.get("_provider", "?")).upper()
    pos = str(r.get("answer_pos", "?"))
    ans = str(r.get("answer", "?"))
    qtype = str(r.get("_q_type", "vocab"))
    qtype_label = "어법" if qtype == "grammar" else "어휘"
    chips = (
        "<span class='chip chip-accent'>" + prov + "</span>"
        "<span class='chip'>" + qtype_label + "</span>"
        "<span class='chip'>" + pos + "</span>"
        "<span class='chip'>정답 " + ans + "번</span>"
    )
    col1, col2 = st.columns([0.9, 0.1])
    with col1:
        st.markdown("<div style='margin-top:1rem'><b>문제 " + str(idx) + "</b> &nbsp; " + chips + "</div>", unsafe_allow_html=True)
        st.markdown("<div class='q-card'>" + r.get("question_text", "") + "</div>", unsafe_allow_html=True)
        with st.expander("정답 · 해설"):
            st.markdown("**정답:** " + ans + "번")
            st.markdown("**원래:** `" + str(r.get("original_word", "?")) + "` → **변형:** `" + str(r.get("modified_word", "?")) + "`")
            st.markdown("**카테고리:** " + pos)
            st.markdown("**해설:** " + str(r.get("explanation", "")))
    with col2:
        if show_delete_button:
            if st.button("🗑️", key=f"del_{r.get('_id', idx)}"):
                delete_question_by_id(r.get("_id"))
                st.rerun()


with tab_q:
    if st.session_state.last_results:
        st.markdown("<div class='divider'></div>", unsafe_allow_html=True)
        st.markdown("### 🧾 이번 세션에서 만든 문제")
        for idx, r in enumerate(st.session_state.last_results, start=1):
            render_question_card(idx, r, show_delete_button=False)

    if user_input.strip():
        saved = questions_for_text(user_input)
        if saved:
            st.markdown("<div class='divider'></div>", unsafe_allow_html=True)
            col_left, col_right = st.columns([1, 1])
            with col_left:
                st.markdown("### 📚 이 지문 누적 문제 (" + str(len(saved)) + "개)")
            with col_right:
                if st.button("🗑️ 이 지문 모든 문제 삭제", key="delete_all_for_text"):
                    delete_questions_for_text(user_input)
                    st.success("현재 지문의 모든 문제가 삭제되었습니다.")
                    st.rerun()
                # Word 내보내기 버튼
                if DOCX_AVAILABLE and saved:
                    docx_buffer = export_questions_to_docx(saved)
                    if docx_buffer:
                        st.download_button(
                            label="📄 Word로 내보내기",
                            data=docx_buffer,
                            file_name=f"word_twist_{datetime.datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                # 인쇄용 HTML (PDF) 버튼
                print_html = """
                <html><head><meta charset="UTF-8"><title>Word Twist 문제</title>
                <style>body { font-family: sans-serif; } .q { margin-bottom: 30px; } u { text-decoration: none; border-bottom: 2px solid #a855f7; }</style>
                </head><body>
                <h1>Word Twist 문제</h1>
                """
                for idx, q in enumerate(saved, 1):
                    print_html += f"""
                    <div class="q">
                    <h3>문제 {idx}</h3>
                    <p>{q.get('question_text', '')}</p>
                    <p><strong>정답:</strong> {q.get('answer', '?')}번</p>
                    <p><strong>원래:</strong> {q.get('original_word', '?')} → <strong>변형:</strong> {q.get('modified_word', '?')}</p>
                    <p><strong>카테고리:</strong> {q.get('answer_pos', '?')}</p>
                    <p><strong>해설:</strong> {q.get('explanation', '')}</p>
                    <hr>
                    </div>
                    """
                print_html += "</body></html>"
                st.download_button(
                    label="🖨️ 인쇄용 HTML (PDF로 저장 가능)",
                    data=print_html,
                    file_name=f"word_twist_print_{datetime.datetime.now().strftime('%Y%m%d_%H%M')}.html",
                    mime="text/html"
                )

            for idx, q in enumerate(saved, start=1):
                render_question_card(idx, q, show_delete_button=True)


# ============= TAB A: 지문 분석 =============
def render_annotated_html(annotated):
    def repl(m):
        full = m.group(1).strip()
        parts = full.split("·", 1)
        comp = parts[0].strip()
        pos = parts[1].strip() if len(parts) > 1 else ""
        comp_cls = "tag-M"
        if comp == "S":
            comp_cls = "tag-S"
        elif comp == "V":
            comp_cls = "tag-V"
        elif comp == "O":
            comp_cls = "tag-O"
        elif comp == "C":
            comp_cls = "tag-C"
        elif comp == "OC":
            comp_cls = "tag-OC"
        chip = "<span class='tag-chip " + comp_cls + "'>" + comp
        if pos:
            chip += "·" + pos
        chip += "</span>"
        return chip
    return re.sub(r"⟦(.*?)⟧", repl, annotated)


with tab_a:
    if not user_input.strip():
        st.info("영어 지문을 입력하면 분석을 시작할 수 있어요.")
    else:
        cached = st.session_state.get("analysis_result")
        cached_for = st.session_state.get("analysis_for_text", "")

        col_a1, col_a2 = st.columns([1, 3])
        with col_a1:
            run_analysis = st.button("📖 지문 분석 시작", key="run_analysis_btn")

        if run_analysis:
            api_keys = {
                "gemini": st.session_state.get("gemini_api_key", ""),
                "openai": st.session_state.get("openai_api_key", ""),
                "anthropic": st.session_state.get("anthropic_api_key", ""),
                "deepseek": st.session_state.get("deepseek_api_key", ""),
            }
            loader = st.empty()
            loader.markdown(
                "<div class='loading-overlay'><div class='loading-box'>"
                "<div class='spinner'></div>"
                "<div class='loading-title'>지문을 분석 중이에요</div>"
                "<div class='loading-sub'>" + provider.upper()
                + " · 문장 단위 직독직해 + 품사 분석"
                + " <span class='loading-dots'></span></div>"
                "<div style='margin-top:14px; color:#94a3b8; font-size:0.78rem'>지문 길이에 따라 15~40초 소요</div>"
                "</div></div>",
                unsafe_allow_html=True,
            )
            try:
                result = analyze_passage(user_input, provider, model, api_keys)
                st.session_state["analysis_result"] = result
                st.session_state["analysis_for_text"] = user_input
                cached = result
                cached_for = user_input
            except Exception as e:
                st.error("분석 실패: " + str(e))
            finally:
                loader.empty()

        if cached and cached_for == user_input:
            sentences = cached.get("sentences", [])
            if not sentences:
                st.warning("분석 결과가 비어있습니다.")
            else:
                st.markdown("<div class='divider'></div>", unsafe_allow_html=True)
                st.markdown("### 📖 분석 결과 (" + str(len(sentences)) + "개 문장)")
                legend = (
                    "<div style='margin-bottom:14px; font-size:0.82rem; color:#9ca3af'>"
                    "<span class='tag-chip tag-S'>S 주어</span> "
                    "<span class='tag-chip tag-V'>V 동사</span> "
                    "<span class='tag-chip tag-O'>O 목적어</span> "
                    "<span class='tag-chip tag-C'>C 보어</span> "
                    "<span class='tag-chip tag-OC'>OC 목적격보어</span> "
                    "<span class='tag-chip tag-M'>M 수식어</span>"
                    "</div>"
                )
                st.markdown(legend, unsafe_allow_html=True)
                for i, s in enumerate(sentences, start=1):
                    annotated_html = render_annotated_html(s.get("annotated", ""))
                    block = (
                        "<div class='sent-block'>"
                        "<div class='sent-num'>문장 " + str(i) + "</div>"
                        "<div class='sent-original'>" + s.get("original", "") + "</div>"
                        "<div class='sent-annotated'>" + annotated_html + "</div>"
                        "<div class='sent-row literal'><span class='lbl'>직독직해</span>"
                        + s.get("literal", "") + "</div>"
                        "<div class='sent-row translation'><span class='lbl'>해석</span>"
                        + s.get("translation", "") + "</div>"
                        "</div>"
                    )
                    st.markdown(block, unsafe_allow_html=True)


# ---------- 드라이브 백업 ----------
st.markdown("<div class='divider'></div>", unsafe_allow_html=True)
cdrv1, cdrv2 = st.columns([1, 3])
with cdrv1:
    drv = st.button("☁️ 드라이브 백업")
with cdrv2:
    st.markdown("<div class='small-dim' style='padding-top:0.5rem'>saved_questions.json 및 passages.json을 구글 드라이브로 업로드합니다.</div>", unsafe_allow_html=True)
if drv:
    found_any = False
    for fname in [SAVE_FILE, PASSAGE_FILE]:
        if os.path.exists(fname):
            found_any = True
            try:
                fid = upload_to_drive(fname)
                if fid:
                    st.success(fname + " 업로드 완료 (ID: " + fid + ")")
            except Exception as e:
                st.error(fname + " 업로드 실패: " + str(e))
    if not found_any:
        st.warning("저장된 파일이 없습니다.")
